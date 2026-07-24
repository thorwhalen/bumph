"""Simple access to docx (Word Doc) elements."""
from io import BytesIO
import docx  # https://automatetheboringstuff.com/chapter13/  (pip install python-docx)
from py2store.util import ModuleNotFoundErrorNiceMessage
from py2store.stores.local_store import LocalBinaryStore
from py2store import wrap_kvs


def get_text_from_docx(doc):
    r"""Get text from a ``docx.Document`` object.

    More precisely, 'text' is the newline-separated concatenation of the
    ``.text`` attributes of every paragraph. You get a document object from a
    file path or pointer ``f`` with ``docx.Document(f)`` (``import docx``;
    ``pip install python-docx``).

    >>> import docx
    >>> doc = docx.Document()
    >>> _ = doc.add_paragraph('hello')
    >>> _ = doc.add_paragraph('world')
    >>> get_text_from_docx(doc)
    'hello\nworld'
    """
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return "\n".join(fullText)


def bytes_to_doc(doc_bytes):
    return docx.Document(BytesIO(doc_bytes))


LocalDocxStore = wrap_kvs(
    LocalBinaryStore, name="LocalDocxStore", obj_of_data=bytes_to_doc
)

LocalDocxTextStore = wrap_kvs(
    LocalDocxStore, name="LocalDocxTextStore", obj_of_data=get_text_from_docx
)
